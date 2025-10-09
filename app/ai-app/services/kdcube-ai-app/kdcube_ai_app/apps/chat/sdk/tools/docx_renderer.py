# SPDX-License-Identifier: MIT
# Copyright (c) 2025 Elena Viter

# chat/sdk/tools/docx_renderer.py

from __future__ import annotations

from pathlib import Path
from typing import List, Dict, Optional, Tuple
import re
import json

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import OxmlElement

from kdcube_ai_app.apps.chat.sdk.runtime.workdir_discovery import resolve_output_dir
import kdcube_ai_app.apps.chat.sdk.tools.md_utils as md_utils

# --------------------------- Helpers / constants -----------------------------

_LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
_CIT_RE  = re.compile(r"\[\[S:(\d+)\]\]")  # [[S:3]]
_CODE_FENCE_RE = re.compile(r"^```(\w+)?\s*$")
_TABLE_ROW_RE  = re.compile(r"^\s*\|.+\|\s*$")
_BLOCKQUOTE_RE = re.compile(r"^\s*>\s+(.*)$")

PALETTE = {
    "fg": RGBColor(20, 24, 31),
    "muted": RGBColor(95, 106, 121),
    "accent": RGBColor(31, 111, 235),
    "quote_bg": RGBColor(245, 247, 250),
    "rule": RGBColor(220, 224, 230),
    "table_header_bg": RGBColor(240, 244, 252),
    "code_bg": RGBColor(250, 250, 252),
}
TYPE = {
    "title": Pt(22),
    "h1": Pt(18),
    "h2": Pt(16),
    "h3": Pt(14),
    "body": Pt(11.5),
    "code": Pt(10.5),
}
MONO = "Consolas"

def _outdir() -> Path:
    return resolve_output_dir()

def _basename_only(path: str, default_ext: str = ".docx") -> str:
    name = Path(path).name
    if default_ext and not name.lower().endswith(default_ext):
        name += default_ext
    return name

def _ensure_parent(p: Path) -> None:
    p.parent.mkdir(parents=True, exist_ok=True)

def _domain_of(url: str) -> str:
    from urllib.parse import urlparse
    try:
        net = urlparse(url).netloc
        return net or url
    except Exception:
        return url

def _add_char_style(run, *, size: Pt, bold=False, italic=False, color: RGBColor | None = None, mono=False):
    font = run.font
    font.size = size
    font.bold = bold
    font.italic = italic
    if mono:
        font.name = MONO
        # East Asian fallback – keeps mono look on Windows
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:eastAsia'), MONO)
    if color is not None:
        font.color.rgb = color

def _set_para(p, *, space_before=3, space_after=3, line_spacing=1.25, align=None):
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    if align:
        p.alignment = align

def _add_heading(doc: Document, text: str, level: int):
    # map to built-in Heading styles
    style = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}.get(level, "Heading 3")
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.clear()  # make sure we control the run
    r = p.add_run(text.strip())
    size = {1: TYPE["h1"], 2: TYPE["h2"], 3: TYPE["h3"]}[level if level in (1,2,3) else 3]
    _add_char_style(r, size=size, bold=True, color=PALETTE["fg"])

def _add_paragraph_text(doc: Document, text: str, level: int = 0):
    # Use Word built-ins for lists to keep bullets nice
    list_style = None
    if re.match(r"^\s*(?:[-*])\s+", text):
        list_style = "List Bullet"
        text = re.sub(r"^\s*[-*]\s+", "", text)
    elif re.match(r"^\s*\d+\.\s+", text):
        list_style = "List Number"
        text = re.sub(r"^\s*\d+\.\s+", "", text)
    p = doc.add_paragraph(style=list_style) if list_style else doc.add_paragraph()
    # indent by level (2 spaces → one level)
    if level > 0:
        p.paragraph_format.left_indent = Inches(0.32 * level)

    _set_para(p, space_before=2, space_after=2)

    # Split for **bold** and *italic*
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2])
            _add_char_style(r, size=TYPE["body"], bold=True, color=PALETTE["fg"])
        elif part.startswith("*") and part.endswith("*"):
            r = p.add_run(part[1:-1])
            _add_char_style(r, size=TYPE["body"], italic=True, color=PALETTE["fg"])
        else:
            _emit_link_or_text(p, part)

    # default any unstyled runs
    for r in p.runs:
        if r.font.size is None:
            _add_char_style(r, size=TYPE["body"], color=PALETTE["fg"])

def _emit_link_or_text(p, text: str, sources_map: Dict[int, Dict[str, str]] = None, resolve_citations: bool = False):
    sources_map = sources_map or {}
    while text:
        m_link = _LINK_RE.search(text)
        m_cit = _CIT_RE.search(text) if resolve_citations else None
        cands = [(m_link, "link"), (m_cit, "cit")]
        cands = [(m, t) for m, t in cands if m]
        if not cands:
            r = p.add_run(text)
            return
        m, kind = min(cands, key=lambda t: t[0].start())
        if m.start() > 0:
            r = p.add_run(text[:m.start()])
        if kind == "link":
            r = p.add_run(m.group(1))
            _add_char_style(r, size=TYPE["body"], color=PALETTE["accent"])
            try:
                r.hyperlink.address = m.group(2)
            except Exception:
                pass
        else:
            sid = int(m.group(1))
            rec = sources_map.get(sid, {})
            label = rec.get("title") or f"[{sid}]"
            url = rec.get("url", "")
            r = p.add_run(label)
            _add_char_style(r, size=TYPE["body"], color=PALETTE["accent"])
            if url:
                try:
                    r.hyperlink.address = url
                except Exception:
                    pass
        text = text[m.end():]

def _add_code_block(doc: Document, lines: List[str]):
    # Use 1x1 table as a card w/ background + border for better reliability
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cell = tbl.cell(0,0)
    # background
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "FAFAFC")  # code_bg
    cell._tc.get_or_add_tcPr().append(shading)
    # border
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for side in ("top","bottom","left","right"):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '6')
        el.set(qn('w:color'), 'DCE0E6')  # rule
        tc_borders.append(el)
    tc_pr.append(tc_borders)

    p = cell.paragraphs[0]
    _set_para(p, space_before=3, space_after=3, line_spacing=1.1)
    for i, ln in enumerate(lines):
        if i > 0:
            p = cell.add_paragraph()
            _set_para(p, space_before=0, space_after=0, line_spacing=1.1)
        r = p.add_run(ln.rstrip("\n"))
        _add_char_style(r, size=TYPE["code"], mono=True, color=PALETTE["fg"])

def _add_blockquote(doc: Document, lines: List[str]):
    # 1x1 table for shaded quote with left rule
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0,0)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F5F7FA")  # quote_bg
    cell._tc.get_or_add_tcPr().append(shading)

    # left rule via table borders
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for side in ("top","bottom","right"):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'nil')
        tc_borders.append(el)
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '16')
    left.set(qn('w:color'), 'DCE0E6')
    tc_borders.append(left)
    tc_pr.append(tc_borders)

    p = cell.paragraphs[0]
    _set_para(p, space_before=2, space_after=2)
    r = p.add_run("\n".join(lines))
    _add_char_style(r, size=TYPE["body"], italic=True, color=PALETTE["muted"])

def _parse_table(block_lines: List[str]) -> Optional[List[List[str]]]:
    rows = [ln.strip() for ln in block_lines if _TABLE_ROW_RE.match(ln)]
    if len(rows) < 2:
        return None
    def split_row(r: str): return [c.strip() for c in r.strip("|").split("|")]
    cells = [split_row(r) for r in rows]
    # second row must be header separator
    if not any(set(c) & {"---", ":---", "---:", ":---:"} for c in cells[1]):
        return None
    hdr = cells[0]
    data = cells[2:] if len(cells) > 2 else []
    return [hdr] + data

def _add_table(doc: Document, data: List[List[str]]):
    rows, cols = len(data), len(data[0])
    tbl = doc.add_table(rows=rows, cols=cols)
    tbl.style = 'Table Grid'
    # header
    for j, txt in enumerate(data[0]):
        cell = tbl.cell(0, j)
        p = cell.paragraphs[0]; p.clear()
        r = p.add_run(txt)
        _add_char_style(r, size=TYPE["body"], bold=True, color=PALETTE["fg"])
        # light header bg
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "F0F4FC")
        cell._tc.get_or_add_tcPr().append(shading)
    # body
    for i in range(1, rows):
        for j, txt in enumerate(data[i]):
            cell = tbl.cell(i, j)
            p = cell.paragraphs[0]; p.clear()
            r = p.add_run(txt)
            _add_char_style(r, size=TYPE["body"], color=PALETTE["fg"])

def _split_markdown_sections(md: str) -> List[Tuple[str, List[str]]]:
    """
    Create sections by '## ' headings; first '# ' becomes doc title if present.
    """
    lines = (md or "").splitlines()
    slides: List[Tuple[str, List[str]]] = []
    cur_title: Optional[str] = None
    cur_body: List[str] = []

    for ln in lines:
        if ln.startswith("## "):
            if cur_title is not None:
                slides.append((cur_title.strip(), cur_body))
            cur_title = ln[3:]
            cur_body = []
        elif ln.startswith("# "):
            if cur_title is None and not slides:
                cur_title = ln[2:]
                cur_body = []
            else:
                cur_body.append(ln)
        else:
            cur_body.append(ln)

    if cur_title is None:
        nonempty = next((l for l in lines if l.strip()), "Document")
        cur_title = nonempty.lstrip("# ").strip() or "Document"

    slides.append((cur_title.strip(), cur_body))
    return slides

# --------------------------- Public entrypoint -------------------------------

def render_docx(
        path: str,
        content_md: str,
        *,
        title: Optional[str] = None,
        sources: Optional[str] = None,
        resolve_citations: bool = True,
        include_sources_section: bool = True
) -> str:
    """
    Render a modern-looking .docx from Markdown with headings, lists, code, quotes, tables, links, citations.
    Returns the **basename** written inside OUTPUT_DIR.
    """
    basename = _basename_only(path, ".docx")
    outdir = _outdir()
    outfile = outdir / basename
    _ensure_parent(outfile)

    sources_map: Dict[int, Dict[str, str]] = {}
    order: List[int] = []
    if sources:
        sources_map, order = md_utils._normalize_sources(sources)

    sections = _split_markdown_sections(content_md or "")
    doc = Document()

    # Title
    title_text = title or sections[0][0]
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_para(p, space_before=0, space_after=4, line_spacing=1.2)
    r = p.add_run(title_text.strip())
    _add_char_style(r, size=TYPE["title"], bold=True, color=PALETTE["fg"])

    # subtle rule
    rule = doc.add_paragraph()
    _set_para(rule, space_before=0, space_after=8)
    r2 = rule.add_run("─" * 40)
    _add_char_style(r2, size=Pt(10), color=PALETTE["rule"])

    # Content (each section = H1 + body)
    for stitle, body in sections:
        # section heading
        _add_heading(doc, stitle, level=1)

        in_code = False
        code_buf: List[str] = []
        table_buf: List[str] = []
        quote_buf: List[str] = []

        def flush_code():
            nonlocal code_buf
            if code_buf:
                _add_code_block(doc, code_buf)
                code_buf = []

        def flush_table():
            nonlocal table_buf
            if table_buf:
                data = _parse_table(table_buf)
                if data:
                    _add_table(doc, data)
                else:
                    for raw in table_buf:
                        _add_paragraph_text(doc, raw)
                table_buf = []

        def flush_quote():
            nonlocal quote_buf
            if quote_buf:
                _add_blockquote(doc, quote_buf)
                quote_buf = []

        for ln in body:
            if _CODE_FENCE_RE.match(ln):
                if not in_code:
                    flush_table(); flush_quote()
                    in_code = True; code_buf = []
                else:
                    flush_code(); in_code = False
                continue
            if in_code:
                code_buf.append(ln)
                continue

            if _TABLE_ROW_RE.match(ln):
                flush_code(); flush_quote()
                table_buf.append(ln)
                continue
            else:
                flush_table()

            m_q = _BLOCKQUOTE_RE.match(ln)
            if m_q:
                flush_code()
                quote_buf.append(m_q.group(1))
                continue
            else:
                flush_quote()

            # Paragraph / list
            # compute indent level from leading spaces
            m_bullet = re.match(r"^(\s*)([-*]|\d+\.)\s+(.*)$", ln)
            if m_bullet:
                spaces, _, tail = m_bullet.groups()
                level = min(len(spaces) // 2, 4)
                _add_paragraph_text(doc, ln, level=level)
            else:
                if ln.strip():
                    _add_paragraph_text(doc, ln)

        # tail flush
        flush_code(); flush_table(); flush_quote()

    if include_sources_section and order:
        _add_heading(doc, "References", level=1)
        for sid in order:
            src = sources_map.get(sid);
            if not src:
                continue
            p = doc.add_paragraph()
            _set_para(p, space_before=1, space_after=1)
            r1 = p.add_run(f"[{sid}] ")
            _add_char_style(r1, size=TYPE["body"], bold=True, color=PALETTE["fg"])
            r2 = p.add_run(src.get("title") or _domain_of(src.get("url","")) or f"Source {sid}")
            _add_char_style(r2, size=TYPE["body"], color=PALETTE["fg"])
            url = src.get("url","")
            if url:
                p2 = doc.add_paragraph()
                _set_para(p2, space_before=0, space_after=6)
                r3 = p2.add_run(url)
                _add_char_style(r3, size=TYPE["body"], color=PALETTE["accent"])

    doc.save(str(outfile))
    return basename
